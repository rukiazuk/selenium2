from curses.ascii import LF
import os
import time
from selenium import webdriver
from selenium.webdriver.common.by import By
import pyperclip
from docx import Document
from selenium.webdriver.common.keys import Keys
from posixpath import expanduser
from selenium.webdriver.support.ui import Select
from PIL import Image
from io import BytesIO
from datetime import date, timedelta
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from extractSpikes import ExtractSpikes
#from extractSpikes import mojascieszka
from selenium.webdriver.common.action_chains import ActionChains


BEG = date.today() - timedelta(days=61)
END = date.today() - timedelta(days=31)
d1 = BEG.strftime("%b %d, %Y")
d2 = END.strftime("%b %d, %Y")
print(d1)
print(d2)

with open("C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\\Desktop\\New folder\\screens\\summary.txt", "w") as f:
    f.write(' ')

SHIPMENTID = pyperclip.paste()
pyperclip.copy(SHIPMENTID)

PATH = "C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\Desktop\\webdriver\\chromedriver.exe"
driver = webdriver.Chrome(PATH)

options = webdriver.ChromeOptions() 
options.add_argument("user-data-dir=C:\\Users\MartaBartkowiak\\OneDrive - Controlant hf\Desktop\webdriver\\newchrome")  #Path to your chrome profile
options.add_argument("download.default_directory=C:/Test")

driver = webdriver.Chrome(executable_path="C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\Desktop\\webdriver\\chromedriver.exe", chrome_options=options)

time.sleep(1)

driver.get("https://scm.controlant.com/global/shipments/search/" + SHIPMENTID + "/chart?page=1")
driver.switch_to.window(driver.current_window_handle)
driver.maximize_window()

time.sleep(3)
#wait for page to load and element appear
try:
    element = WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.XPATH, "//label[normalize-space()='Light events']")))
except:
    driver.quit()
#time.sleep(10)

driver.find_element(By.XPATH, "//label[normalize-space()='Light events']").click()

time.sleep(5)

driver.find_element(By.XPATH, "//button[@class='_multiSelectButton_11xmn_7']").click()
LOGGERID1 = driver.find_element(By.XPATH, "//label[@class='_multiSelectItem_11xmn_59']").text 

image_chart = driver.find_element(By.CSS_SELECTOR, "#single-shipment-view > div._verticalTop_hzerk_18._hideScroll_hzerk_27 > div > div._pane_hzerk_77._chartPane_hzerk_94 > div").screenshot("C:\\Users\MartaBartkowiak\\OneDrive - Controlant hf\\Desktop\\New folder\\screens\\test.png")

time.sleep(2)
driver.find_element(By.XPATH, "//a[@data-testid='anchor-data']").click() 
time.sleep(2)
driver.find_element(By.XPATH, "//button[normalize-space()='Download CSV']").click() 

driver.find_element(By.XPATH, "//a[@data-testid='anchor-info']").click()

time.sleep(5)

#data exctract
SHIPPED = driver.find_element(By.CSS_SELECTOR, "div[data-testid='span-shipment-shipped']").text 

ch = "2023"    
stripped1 = SHIPPED.strip("Shipped")
head, sep, tail = stripped1.partition(ch)
DATEBEG = head + sep
print(DATEBEG)
DELIVERED = driver.find_element(By.CSS_SELECTOR, "div[data-testid='span-shipment-delivered']").text

ch = "2023"
stripped2 = DELIVERED.strip("Delivered")
head, sep, tail = stripped2.partition(ch)
DATEEND = head + sep
print(DATEEND)

OBD = driver.find_element(By.CSS_SELECTOR, "div[class='_shipmentName_hzerk_54'] span").text
print(OBD)
ORIGIN = driver.find_element(By.CSS_SELECTOR, "span[data-testid='shipment-origin-name'] strong").text
DESTINATION = driver.find_element(By.CSS_SELECTOR, "span[data-testid='shipment-destination-name'] strong").text

#hover and get data missing points
try:
    a = ActionChains(driver)
    m = driver.find_element(By.XPATH, "//div[@class='_loggerIdContainer_1m880_22']//*[name()='svg']//*[name()='path' and contains(@fill,'currentCol')]") 
    
    a.move_to_element(m).perform()
    elements = driver.find_elements_by_xpath("/html/body/div/div[2]/div[4]/div/main/div/div[1]/div/div[3]/div/main/div[3]/div[2]/div/div[2]/div/div[2]/div/div/div/div[2]")
    for e in elements:
        print(e.text)

    with open("C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\\Desktop\\New folder\\screens\\summary.txt", "w") as f:
        for e in elements:
            f.write(e.text)
            f.write('\n')
except:
    print("no data missing")

time.sleep(2)
driver.find_element(By.CSS_SELECTOR, "._loggerLink_1m880_64").click()
time.sleep(3)
driver.find_element(By.XPATH, "//a[@data-testid='info-tab-button']").click()
time.sleep(3)
FWVERSION_RAW = driver.find_element(By.XPATH, "/html/body/div/div[2]/div[4]/div/main/div/div[1]/div/div[2]/div/div[4]/p[1]").text
FWVERSION = FWVERSION_RAW.replace("Firmware version", "Proc")
#FWVERSION = driver.find_element(By.XPATH, "/html/body/div/div[2]/div[4]/div/main/div/div[1]/div/div[2]/div/div[4]/p[1]/text()").text
LASTCHIN = driver.find_element(By.CSS_SELECTOR, "#single-logger > div._verticalTop_4zxoj_33 > div > div._pane_4zxoj_65 > div > div._loggerInformation_1tjqt_56 > p._lastCheckIn_1tjqt_130 > strong").text
LOGGERID = driver.find_element(By.CSS_SELECTOR, "#single-logger > div._verticalTop_4zxoj_33 > div > div._loggerHeader_4zxoj_1 > span").text

ORIDEN = str(ORIGIN + " / " + DESTINATION)

time.sleep(5)

# >>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>> second part


#PATH = "C:\\Users\\MartaBartkowiak\\OneDrive - C
# ontrolant hf\Desktop\\webdriver\\chromedriver.exe"
#driver = webdriver.Chrome(PATH)

#options = webdriver.ChromeOptions() 
#options.add_argument("user-data-dir=C:\\Users\MartaBartkowiak\\OneDrive - Controlant hf\Desktop\webdriver\\newchrome2")  #Path to your chrome profile

#driver = webdriver.Chrome(executable_path="C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\Desktop\\webdriver\\chromedriver.exe", chrome_options=options)

driver.get("https://reseller.controlant.com/search/" + LOGGERID + "/device-info")
driver.maximize_window()

time.sleep(2)


try:
    element = WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.XPATH, "//h1[1]")))
except:
    driver.get("https://reseller.controlant.com/search/" + LOGGERID + "/device-info")

time.sleep(5)
driver.get("https://reseller.controlant.com/search/" + LOGGERID + "/device-info")
#wait for page to load and element appear
#try:
#    element = WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.XPATH, By.XPATH, "/html[1]/body[1]/div[1]/main[1]/div[1]/section[1]/div[1]/div[1]/table[1]/tbody[1]/tr[11]/td[1]/input[1]")))
#except:
#    driver.quit()

#time.sleep(5)

#set the date
def SetTheDate():

    time.sleep(4)
    driver.find_element(By.XPATH, "/html/body/div[1]/main/div/section/div[1]/div[1]/table/tbody/tr[11]/td[1]/input").click() 
    driver.find_element(By.XPATH, "//div[@view='chart data']//select[@class='ng-pristine ng-untouched ng-valid ng-not-empty']").click()

    select = Select(driver.find_element(By.XPATH, "//div[@view='chart data']//select[@class='ng-pristine ng-untouched ng-valid ng-not-empty']"))
    select.select_by_index(2)

    driver.find_element(By.XPATH, "//div[@class='customPeriod']//input[@id='datepicker-start']").click()
    driver.find_element(By.XPATH, "//div[@class='customPeriod']//input[@id='datepicker-start']").clear()
    driver.find_element(By.XPATH, "//div[@class='customPeriod']//input[@id='datepicker-start']").send_keys(DATEBEG)
    driver.find_element(By.XPATH, "//form[@class='ng-valid ng-dirty ng-valid-parse']//input[@id='datepicker-end']").click()
    driver.find_element(By.XPATH, "//form[@class='ng-valid ng-dirty ng-valid-parse']//input[@id='datepicker-end']").clear()
    driver.find_element(By.XPATH, "//form[@class='ng-valid ng-dirty ng-valid-parse']//input[@id='datepicker-end']").send_keys(DATEEND)


#change date -/+ day
def ChangeDay():
    time.sleep(1)
    driver.find_element(By.XPATH, "//form[@class='ng-valid ng-dirty ng-valid-parse']//input[@id='datepicker-start']").send_keys(Keys.ARROW_LEFT)
    driver.find_element(By.XPATH, "//form[@class='ng-valid ng-dirty ng-valid-parse']//label[@for='datepicker-end'][normalize-space()='To:']").send_keys(Keys.ARROW_RIGHT + Keys.ARROW_RIGHT)
    driver.find_element(By.XPATH, "//form[@class='ng-valid ng-dirty ng-valid-parse']//button[@class='small'][normalize-space()='Apply']").click()


#take screenshot and crop
def Screen():
    time.sleep(10)
    LTScreen = driver.get_screenshot_as_png()
    LTCropped = Image.open(BytesIO(LTScreen))
    #LTCropped = LTCropped.crop((20, 85, 2360, 940))    #everything, but too big
    LTCropped = LTCropped.crop((20, 85, 2360, 740))     #LT only
    path2 = expanduser("C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\\Desktop\\New folder\\screens\\")
    LTCropped.save(path2 + "LT Screen.png")


#LT screens for last 2 months
def LT1ScreenMonths1():
    select = Select(driver.find_element(By.XPATH, "//select[@class='ng-valid ng-not-empty ng-dirty ng-valid-parse ng-touched']"))
    select.select_by_index(1)
    time.sleep(1)
    driver.find_element(By.XPATH, "//h1[1]").click()
    time.sleep(5)
    LT1Screen = driver.get_screenshot_as_png()
    LT1Cropped = Image.open(BytesIO(LT1Screen))
    LT1Cropped = LT1Cropped.crop((20, 85, 2360, 740))     #LT only
    path3 = expanduser("C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\\Desktop\\New folder\\screens\\")
    LT1Cropped.save(path3 + "LT1.png")

def LT2ScreenMonths2():
    time.sleep(2)
    select = Select(driver.find_element(By.XPATH, "//select[@class='ng-valid ng-not-empty ng-dirty ng-valid-parse ng-touched']"))
    select.select_by_index(2)

    driver.find_element(By.XPATH, "//div[@class='customPeriod']//input[@id='datepicker-start']").click()
    driver.find_element(By.XPATH, "//div[@class='customPeriod']//input[@id='datepicker-start']").clear()
    driver.find_element(By.XPATH, "//div[@class='customPeriod']//input[@id='datepicker-start']").send_keys(d1)

    time.sleep(2)

    try:
        element = WebDriverWait(driver, 3).until(EC.presence_of_element_located((By.XPATH, "//form[@class='ng-valid ng-dirty ng-submitted ng-valid-parse']//input[@id='datepicker-end']")))
        time.sleep(2)
        driver.find_element(By.XPATH, "//form[@class='ng-valid ng-dirty ng-submitted ng-valid-parse']//input[@id='datepicker-end']").click()
        driver.find_element(By.XPATH, "//form[@class='ng-valid ng-dirty ng-submitted ng-valid-parse']//input[@id='datepicker-end']").clear()
        driver.find_element(By.XPATH, "//form[@class='ng-valid ng-dirty ng-submitted ng-valid-parse']//input[@id='datepicker-end']").send_keys(d2)
        time.sleep(2)
        driver.find_element(By.XPATH, "//h1[1]").click()
        driver.find_element(By.XPATH, "//form[@class='ng-valid ng-dirty ng-submitted ng-valid-parse']//button[@class='small'][normalize-space()='Apply']").click()
    except:
        driver.find_element(By.XPATH, "//form[@class='ng-valid ng-dirty ng-valid-parse ng-submitted']//input[@id='datepicker-end']").click()
        driver.find_element(By.XPATH, "//form[@class='ng-valid ng-dirty ng-valid-parse ng-submitted']//input[@id='datepicker-end']").clear()
        driver.find_element(By.XPATH, "//form[@class='ng-valid ng-dirty ng-valid-parse ng-submitted']//input[@id='datepicker-end']").send_keys(d2)
        time.sleep(2)
        driver.find_element(By.XPATH, "//h1[1]").click()
        driver.find_element(By.XPATH, "//form[@class='ng-valid ng-dirty ng-valid-parse ng-submitted']//button[@class='small'][normalize-space()='Apply']").click()

   
    
   # time.sleep(2)
   # driver.find_element(By.XPATH, "//h1[1]").click()
   # driver.find_element(By.XPATH, "//form[@class='ng-valid ng-dirty ng-submitted ng-valid-parse']//button[@class='small'][normalize-space()='Apply']").click()
    #N = 30  # number of times you want to press key
    #actions = ActionChains(driver) 
    #for _ in range(N):
    #    actions = actions.send_keys(Keys.ARROW_LEFT)
    #actions.perform()

    time.sleep(3)

    LT2Screen = driver.get_screenshot_as_png()
    LT2Cropped = Image.open(BytesIO(LT2Screen))
    LT2Cropped = LT2Cropped.crop((20, 85, 2360, 740))     #LT only
    path4 = expanduser("C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\\Desktop\\New folder\\screens\\")
    LT2Cropped.save(path4 + "LT2.png")

#generate file
def GenerateDocxFile():
    document = Document()

    document.add_paragraph(LASTCHIN)
    document.add_paragraph(OBD)
    document.add_paragraph(ORIGIN + " / " + DESTINATION)
    
    #document.add_paragraph(SHIPPED)
    #document.add_paragraph(DELIVERED)
    document.add_paragraph(" ")
    document.add_paragraph("Data missing:")
    document.add_paragraph(" ")

    #document.add_picture("C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\\Desktop\\New folder\\screens\\scm.png")

    document.add_paragraph("Logger temperature within normal boundaries. / Modem temperature was below operational boundaries.")
    
    #document.add_picture("C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\\Desktop\\New folder\\screens\\LT Screen.png")

    document.add_paragraph("No probe failure events displayed.")

    document.save("C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\\Desktop\\New folder\\screens\\screensdemo.docx")


def FindPIEvents():
    select2 = Select(driver.find_element(By.XPATH, "/html/body/div[1]/main/div/section/div[2]/div[1]/label/select"))
    select2.select_by_index(1)
    time.sleep(1)
    #driver.find_element(By.XPATH, "//h1[1]").click()
    try:
        time.sleep(1)
        driver.find_element(By.XPATH, "//td[@class='field event-id ng-binding'][normalize-space()='PI']").click()
        time.sleep(10)
        #PIScreen = driver.get_screenshot_as_png()
        #PICropped = Image.open(BytesIO(PIScreen))
        #PICropped = PICropped.crop((1440, 2000, 2300, 2150))     #PI only
        #path3 = expanduser("C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\\Desktop\\New folder\\screens\\")
        #PICropped.save(path3 + "PI.png")
    except:
        print("no PI")


lines = [FWVERSION, LASTCHIN, OBD, ORIGIN + " / " + DESTINATION, " ", "Data missing: ", "Logger temperature within normal boundaries. / Logger temperature was below operational boundaries.", "No probe issue events displayed. / Probe issue event displayed."]

with open("C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\\Desktop\\New folder\\screens\\summary.txt", "a") as f:
    for line in lines:
        f.write(line)
        f.write('\n')

SetTheDate()
ChangeDay()
Screen()

LT1ScreenMonths1()
LT2ScreenMonths2()

FindPIEvents()

with open("C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\\Desktop\\New folder\\screens\\summary.txt") as f:
    data = f.read()
    pyperclip.copy(data)



# >>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>> third part

driver.get("https://controlant.atlassian.net/wiki/spaces/PROD/pages/1212940318/COPY+Summary+Report+-+300445C5+-+2022-11-24")
driver.switch_to.window(driver.current_window_handle)
driver.maximize_window()

time.sleep(2)


try:
    element = WebDriverWait(driver, 130).until(EC.presence_of_element_located((By.XPATH, "//a[@id='editPageLink']//span[@role='presentation']")))
except:
    driver.quit()

time.sleep(4)
driver.find_element(By.XPATH, "//a[@id='editPageLink']//span[@role='presentation']").click()

time.sleep(10)



def LFRgenerate():
#paste everything
    tit = driver.find_element(By.XPATH, "//textarea[@placeholder='Page title']")
    tit.send_keys(Keys.CONTROL + "a")
    tit.send_keys("COPY Summary Report - " + LOGGERID + " - 2023-01-XX")
    tit.send_keys(Keys.TAB*6 + Keys.ENTER + "RH-2331 \\ RH-2332" + Keys.TAB*2 + LOGGERID + Keys.TAB*6 + "Modem: 1.0.0" + Keys.ENTER + "PROC: " + FWVERSION)
    time.sleep(5)

    t = driver.find_element(By.XPATH, "/html[1]/body[1]/div[1]/div[2]/div[1]/div[1]/div[1]/div[1]/div[2]/div[1]/div[2]/div[3]/div[2]/div[2]/div[1]/div[1]/div[1]/div[1]/div[4]/div[1]/span[1]/div[1]/div[1]/div[3]/div[1]/div[1]/div[1]/div[1]/div[3]/table[1]/tbody[1]/tr[1]/td[1]/p[1]")
    
    t.click()
    t.send_keys(Keys.TAB*15 + LASTCHIN + Keys.TAB*2 + OBD + Keys.TAB*2 + ORIDEN + Keys.TAB*4 + "Data missing in the shipment" + Keys.TAB*2 + "N/A" + Keys.TAB*2 + "Data_Missing" + Keys.TAB*2 + "data z pyperclip + screeny") 
    time.sleep(3)
    t.click()
    time.sleep(1)
    t.send_keys(Keys.TAB*31 + "Device error" + Keys.ENTER + "The device cannot connect to the probe temperature sensor in every measuring point") 
    time.sleep(3)
    t.click()
    t.send_keys(Keys.TAB*33 + "Recall the device and the probe")
    time.sleep(3)
    t.click()
    t.send_keys(Keys.TAB*37 + "@")



def LFRtaketwo():
    actions3 = ActionChains(driver)
    actions4 = ActionChains(driver)
    actions5 = ActionChains(driver)
    actions3.send_keys(Keys.TAB*5 + Keys.ENTER + "XXXXXX" + Keys.TAB*2 + LOGGERID  + Keys.TAB*6 + "Modem: 1.0.0" + Keys.ENTER + FWVERSION + Keys.ARROW_UP*2 + Keys.TAB*4 + LASTCHIN + Keys.TAB*2 + OBD + Keys.TAB*2 + ORIDEN + Keys.TAB*4 + "Data missing in the shipment" + Keys.TAB*2 + "N/A" + Keys.TAB*2 + "Data_Missing")
    actions4.send_keys(Keys.TAB*2 + "data z pyperclip + screeny" + Keys.TAB*2 + "Device error" + Keys.ENTER + "The device cannot connect to the probe temperature sensor in every measuring point" + Keys.ARROW_UP*3) 
    actions5.send_keys(Keys.TAB*4 + "Recall the device and the probe" + Keys.TAB*4 + "@" )
    actions3.perform()
    #time.sleep(2)
    #actions4.perform()
    #time.sleep(2)
    #actions5.perform()

    tit = driver.find_element(By.XPATH, "/html[1]/body[1]/div[1]/div[2]/div[1]/div[1]/div[1]/div[1]/div[2]/div[1]/div[2]/div[3]/div[2]/div[2]/div[1]/div[1]/div[1]/div[1]/div[2]/div[2]/div[2]/textarea[1]")
    tit.send_keys(Keys.CONTROL + "a")
    tit.send_keys("COPY Summary Report - " + LOGGERID + " - 2023-01-XX")

#paste 
#time.sleep(2)
#driver.find_element(By.XPATH, "//span[contains(text(),'Publish')]").click()
time.sleep(2)
#LFRgenerate()
LFRtaketwo()

time.sleep(10)

img1 = Image.open("C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\\Desktop\\New folder\\screens\\test.png")
img2 = Image.open("C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\\Desktop\\New folder\\screens\\LT Screen.png")
img3 = Image.open("C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\\Desktop\\New folder\\screens\\LT1.png")
img4 = Image.open("C:\\Users\\MartaBartkowiak\\OneDrive - Controlant hf\\Desktop\\New folder\\screens\\LT2.png")
img1.show()
img2.show()
img3.show()
img4.show()

driver.switch_to.window(driver.current_window_handle)
ExtractSpikes()

## If file exists, delete it ##
#if os.path.isfile(mojascieszka):
    #os.remove(mojascieszka)
#else:    ## Show an error ##
    #print("Error: %s file not found" % mojascieszka)


 #   //form[@class='ng-valid ng-dirty ng-valid-parse ng-submitted']//input[@id='datepicker-end']